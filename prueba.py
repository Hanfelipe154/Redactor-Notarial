import streamlit as st
from docx import Document
from datetime import date
from campos_por_documento import documentos
import io
import os
import re

st.set_page_config(page_title="Redactor Notarial", layout="centered")
st.title("Redactor Inteligente para Documentos Notariales")

# Selección del tipo de documento
tipo_doc = st.selectbox("Selecciona el tipo de documento:", list(documentos.keys()))
plantilla_path = documentos[tipo_doc]["plantilla"]
campos = documentos[tipo_doc]["campos"]

from departamentos_ciudades import departamentos_ciudades

departamentos_colombia = list(departamentos_ciudades.keys())

selectbox_personalizados = {
    "estadocivil_poderdante": ["Soltero(a)", "Casado(a)", "Unión marital de hecho", "Divorciado(a)", "Viudo(a)"],
    "estadocivil_poderdante1": ["Soltero(a)", "Casado(a)", "Unión marital de hecho", "Divorciado(a)", "Viudo(a)"],
    "estadocivil_poderdante2": ["Soltero(a)", "Casado(a)", "Unión marital de hecho", "Divorciado(a)", "Viudo(a)"],
    "afectainmueble": ["Sí", "No"],
    "actos": [
        "Compra", "Venta", "venta nuda propiedad y reserva de usufructo", "compra nuda propiedad y reserva de usufructo", "Cancelación de hipoteca", 
        "cancelación de afectación a vivienda familiar", "cancelación de patrimonio de familia", "cancelacion de condicion resolutoria y prohibición de transferencia",
        "Constitución de usufructo",
        "Afectación a vivienda familiar", "Hipoteca", "Levantamiento de patrimonio"
    ]
}

# Inicializa el diccionario
datos = {}
st.subheader("✍️ Complete los campos:")

# Manejo especial de deparinmueble y ciudadinmueble
if "deparinmueble" in campos:
    datos["deparinmueble"] = st.selectbox("Departamento donde se Ubica el Inmueble", departamentos_colombia)
else:
    datos["deparinmueble"] = ""

if "ciudadinmueble" in campos:
    ciudades = departamentos_ciudades.get(datos["deparinmueble"], []) if datos["deparinmueble"] else []
    datos["ciudadinmueble"] = st.selectbox("Ciudad donde se Ubica el Inmueble", ciudades)
else:
    datos["ciudadinmueble"] = ""

# Resto de campos dinámicos
for campo in campos:
    if campo in ["deparinmueble", "ciudadinmueble"]:
        continue  # Ya fueron tratados

    label = campo.replace("_", " ").capitalize()

    if campo in selectbox_personalizados:
        datos[campo] = st.selectbox(label, selectbox_personalizados[campo])
    else:
        datos[campo] = st.text_input(label)

def renderizar_docx_avanzado(path_docx, datos):
    doc = Document(path_docx)

    def reemplazar_en_parrafo(p):
        texto_completo = "".join(run.text for run in p.runs)
        for clave, valor in datos.items():
            texto_completo = texto_completo.replace(f"{{{{{clave}}}}}", str(valor))

        # Vacía todos los runs y escribe uno nuevo con el texto completo
        if p.runs:
            p.runs[0].text = texto_completo
            for i in range(1, len(p.runs)):
                p.runs[i].text = ""

    # Reemplazo en párrafos
    for p in doc.paragraphs:
        reemplazar_en_parrafo(p)

    # Reemplazo en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    reemplazar_en_parrafo(p)

    return doc

# Al presionar el botón
if st.button("Generar documento"):
    
    doc_temp = Document(plantilla_path)
   
    def detectar_variables(doc):
        variables = set()

        for p in doc.paragraphs:
            texto_completo = "".join([r.text for r in p.runs])
            matches = re.findall(r"{{(.*?)}}", texto_completo)
            variables.update([m.strip() for m in matches])

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        texto_completo = "".join([r.text for r in p.runs])
                        matches = re.findall(r"{{(.*?)}}", texto_completo)
                        variables.update([m.strip() for m in matches])

        return variables

    variables_encontradas = detectar_variables(doc_temp)
    st.subheader("🔍 Variables detectadas en la plantilla:")
    st.write(variables_encontradas)

    variables_ingresadas = set(datos.keys())
    faltantes = variables_encontradas - variables_ingresadas

    if faltantes:
        st.warning(f"⚠️ Faltan datos para estas variables: {faltantes}")
        st.stop()  # Detiene la ejecución para que no genere el archivo
    else:
        try:
            doc_generado = renderizar_docx_avanzado(plantilla_path, datos)

            buffer = io.BytesIO()
            doc_generado.save(buffer)
            buffer.seek(0)

            nombres_posibles = ["poderdante", "poderdante1"]
            nombre_persona = next((datos[c] for c in nombres_posibles if c in datos and datos[c]), "documento")
            nombre_persona = nombre_persona.strip().lower().replace(" ", "_")
            nombre_archivo = f"{tipo_doc.lower().replace(' ', '_')}_{nombre_persona}.docx"

            st.success("✅ Documento generado exitosamente.")
            st.download_button(
                label="📥 Descargar documento",
                data=buffer,
                file_name=nombre_archivo,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        except Exception as e:
            st.error(f"❌ Error al generar el documento: {e}")
